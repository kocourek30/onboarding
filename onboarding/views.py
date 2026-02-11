from datetime import date
import os
from io import BytesIO

from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.files.base import ContentFile
from django.db import models
from django.http import Http404, HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse_lazy
from django.utils.encoding import smart_str
from django.views.generic import ListView, CreateView
from docx import Document

from provozy.models import Provoz
from .forms import OsobniDotaznikForm
from .models import OsobniDotaznik, Pozice
from docx2pdf import convert
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages

# -----------------------
# DOCX – pracovní smlouva
# -----------------------

@login_required
def smlouva_pracovni_pomer_docx(request, pk):
    obj = get_object_or_404(OsobniDotaznik, pk=pk)
    user = request.user
    if not user.is_superuser and getattr(user, "role", None) != user.HR:
        if obj.provoz not in user.provozy.all():
            raise Http404()

    template_path = os.path.join(
        settings.BASE_DIR, "templates", "docx", "smlouva_pracovni_pomer.docx"
    )
    document = Document(template_path)

    provoz = obj.provoz
    client_name = provoz.nazev
    client_address_parts = [
        provoz.ulice,
        provoz.mesto,
        provoz.psc,
    ]
    client_address = ", ".join([p for p in client_address_parts if p])

    ctx = {
        "first_name": obj.jmeno,
        "last_name": obj.prijmeni,
        "perm_street": obj.trv_ulice,
        "perm_house_number": obj.trv_cislo,
        "perm_city": obj.trv_mesto,
        "perm_zip": obj.trv_psc,
        "date_of_birth": obj.datum_narozeni.strftime("%d.%m.%Y") if obj.datum_narozeni else "",
        "start_date": obj.nastup_datum.strftime("%d.%m.%Y") if obj.nastup_datum else "",
        "position_name": smart_str(obj.pozice.nazev),
        "workplace_address": f"{client_name}, {client_address}" if client_address else client_name,
        "operation_number": str(provoz.cislo_provozu),

        # klient (název + adresa)
        "client": f"{client_name}, {client_address}" if client_address else client_name,

        "contract_type": (
            "na dobu neurčitou" if obj.typ_pomeru == "DOBA_NEURCITA" else "na dobu určitou"
        ),
        "contract_fixed_until": obj.pomer_do.strftime("%d.%m.%Y") if obj.pomer_do else "",
        "trial_period_months": str(obj.zkusebni_doba_mesice or ""),
        "weekly_hours": str(obj.tydenni_uvazek_hodin or ""),
        "worktime_layout": (
            "rovnoměrně" if obj.rozvrzeni_pracovni_doby == "ROVNOMERNE" else "nerovnoměrně"
        ),
    }

    def replace_text_in_paragraph(paragraph, mapping):
        for key, value in mapping.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    for paragraph in document.paragraphs:
        replace_text_in_paragraph(paragraph, ctx)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, ctx)

    buf = BytesIO()
    document.save(buf)
    buf.seek(0)

    filename_docx = f"smlouva_{obj.prijmeni}_{obj.jmeno}.docx"

    # uložit DOCX do FileFieldu (přepíše starší verzi)
    if hasattr(obj, "smlouva_pracovni_pomer"):
        obj.smlouva_pracovni_pomer.save(
            filename_docx,
            ContentFile(buf.getvalue()),
            save=True,
        )

    # pokus o vytvoření PDF z uloženého DOCX
    if getattr(obj, "smlouva_pracovni_pomer", None):
        docx_path = obj.smlouva_pracovni_pomer.path
        pdf_filename = filename_docx.replace(".docx", ".pdf")
        pdf_rel_path = os.path.join("smlouvy", pdf_filename)
        pdf_abs_path = os.path.join(settings.MEDIA_ROOT, pdf_rel_path)

        os.makedirs(os.path.dirname(pdf_abs_path), exist_ok=True)

        try:
            convert(docx_path, pdf_abs_path)
            if hasattr(obj, "smlouva_pracovni_pomer_pdf"):
                obj.smlouva_pracovni_pomer_pdf.name = pdf_rel_path
                obj.save(update_fields=["smlouva_pracovni_pomer_pdf"])
        except Exception:
            # když převod selže, PDF prostě nebude
            pass

    # vrátit DOCX jako download
    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename_docx}"'
    return response



# -----------------------
# Jednokrokový dotazník
# -----------------------

@login_required
def dotaznik_list(request):
    qs = OsobniDotaznik.objects.select_related("provoz", "pozice", "vytvoril")
    user = request.user

    # superuser a HR vidí vše, ostatní jen svoje provozy
    if not user.is_superuser and getattr(user, "role", None) != user.HR:
        qs = qs.filter(provoz__in=user.provozy.all())

    provoz_id = request.GET.get("provoz") or ""
    pozice_id = request.GET.get("pozice") or ""
    q = request.GET.get("q") or ""

    if provoz_id:
        qs = qs.filter(provoz_id=provoz_id)
    if pozice_id:
        qs = qs.filter(pozice_id=pozice_id)
    if q:
        qs = qs.filter(
            models.Q(jmeno__icontains=q)
            | models.Q(prijmeni__icontains=q)
            | models.Q(provoz__nazev__icontains=q)
            | models.Q(pozice__nazev__icontains=q)
        )

    provozy = Provoz.objects.all()
    pozice = Pozice.objects.all()

    context = {
        "dotazniky": qs,
        "provozy": provozy,
        "pozice_list": pozice,
        "filtr_provoz": provoz_id,
        "filtr_pozice": pozice_id,
        "filtr_q": q,
    }
    return render(request, "onboarding/dotaznik_list.html", context)


@login_required
def dotaznik_create(request):
    if request.method == "POST":
        form = OsobniDotaznikForm(request.POST, user=request.user)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.vytvoril = request.user
            obj.save()
            return redirect("dotaznik_list")
    else:
        form = OsobniDotaznikForm(user=request.user)

    return render(request, "onboarding/dotaznik_form.html", {"form": form})


@login_required
def dotaznik_detail(request, pk):
    obj = get_object_or_404(OsobniDotaznik, pk=pk)
    user = request.user

    if not user.is_superuser and getattr(user, "role", None) != user.HR:
        if obj.provoz not in user.provozy.all():
            raise Http404()

    return render(request, "onboarding/dotaznik_detail.html", {"obj": obj})


@login_required
def dotaznik_update(request, pk):
    obj = get_object_or_404(OsobniDotaznik, pk=pk)
    user = request.user

    if not user.is_superuser and getattr(user, "role", None) != user.HR:
        if obj.provoz not in user.provozy.all():
            raise Http404()

    if request.method == "POST":
        form = OsobniDotaznikForm(request.POST, instance=obj, user=request.user)
        if form.is_valid():
            form.save()
            return redirect("dotaznik_detail", pk=obj.pk)
    else:
        form = OsobniDotaznikForm(instance=obj, user=request.user)

    return render(request, "onboarding/dotaznik_form.html", {"form": form, "obj": obj})


# -----------------------
# Moje provozy
# -----------------------

class MojeProvozyView(LoginRequiredMixin, ListView):
    model = Provoz
    template_name = "onboarding/moje_provozy.html"
    context_object_name = "provozy"

    def get_queryset(self):
        user = self.request.user
        if user.is_superuser or getattr(user, "role", None) == user.HR:
            return Provoz.objects.all()
        return Provoz.objects.filter(uzivatele=user).distinct()


# -----------------------
# (volitelné) CBV varianta create
# -----------------------

class OsobniDotaznikCreateView(LoginRequiredMixin, CreateView):
    model = OsobniDotaznik
    form_class = OsobniDotaznikForm
    template_name = "onboarding/dotaznik_form.html"

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["user"] = self.request.user
        return kwargs

    def form_valid(self, form):
        form.instance.vytvoril = self.request.user
        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy("dotaznik_detail", args=[self.object.pk])


@login_required
def dotaznik_delete(request, pk):
    dotaznik = get_object_or_404(OsobniDotaznik, pk=pk, vytvoril=request.user)
    if request.method == "POST":
        dotaznik.delete()
        messages.success(request, "Dotazník byl odstraněn.")
        return redirect("dotaznik_list")
    return redirect("dotaznik_list")